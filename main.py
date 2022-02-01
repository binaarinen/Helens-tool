#!/venv/bin/python
#
# Exampleapplication Helen
# version 0
# (c) 2021 T.Wisotzki
#

import os.path
import tkinter as tk
import tkinter.messagebox as msg
from docx import Document


# global variable - defines where text-files are stored
filepath = "texts/"


class DocumentCreator(tk.Tk):
    def __init__(self, optionnames):
        super().__init__()
        self.title("Document creator")
        self.geometry("300x400")

        self.options = []
        self.choosenOptions = []

        # instructions at top
        label = tk.Label(self, text="Select textblocks to include:")
        label.pack()

        # create checkboxes for every file
        for opname in optionnames:
            newOption = tk.Checkbutton(self, text=opname)
            newOption.bind("<Button-1>", self.add_option)
            self.options.append(newOption)
        for option in self.options:
            option.pack()

        # generation Button
        self.generateButton = tk.Button(
            self, text="generate document", command=self.generateWord
        )
        self.generateButton.pack()

        self.exitButton = tk.Button(self, text="exit", command=self.exitApp)
        self.exitButton.pack()

    # if text-part is de-/selected it is added/removed
    def add_option(self, event):
        task = event.widget.cget("text")
        if task in self.choosenOptions:
            self.choosenOptions.remove(task)
        else:
            self.choosenOptions.append(task)

    def exitApp(self):
        exit()

    #    # generation of combined textfile
    #    def generate(self):
    #        print('generating file...')
    #        self.generateWord()
    #        with open('final.txt', 'w') as f:
    #            for op in self.choosenOptions:
    #                with open(filepath + op + '.txt', 'r') as r:
    #                    f.write(r.read())

    # generation of combined file as word
    def generateWord(self):
        document = Document()
        document.add_heading("Example Document template")
        for op in self.choosenOptions:
            with open(filepath + op + ".txt", "r") as r:
                p = document.add_paragraph("")
                p.add_run(r.read())
        document.save("final.docx")


# get all text-blocks in specified directory
def getOptions():
    if not os.path.isdir(filepath):
        print("Error: path to text extracts not valid")
        exit()
    options = list()
    for option in os.listdir(filepath):
        options.append(option.split(".")[0])
    return options


if __name__ == "__main__":
    documentcreator = DocumentCreator(getOptions())
    documentcreator.mainloop()
